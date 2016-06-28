''' Python code run for RFS folder
'''


import os
import docx
import re

cur = os.getcwd()  # current working directory

result = open('Result_RFS.txt', 'wb')

rfs = cur + '\\RFS'  
os.chdir(rfs)

ncur = os.getcwd()  # new current working directory

files = os.listdir('.') # files under current working directory

journal = 'The Review of Financial Studies'



# header
header = "Year\tTitle\tJournal"
for i in range(1,16):
	header = header + '\tJEL Code ' + str(i)
header = header + '\n'
result.write(header)

for f in files:

	year = f[0:4]
	title = f[5:-5]

	path =  f
	doc = docx.Document(path)

	pattern = re.compile('\(JEL.+\)')

	for i in range(3, len(doc.paragraphs)):
		text = doc.paragraphs[i].text
		matches = pattern.findall(text)
		if len(matches) == 0:
			continue
		elif len(matches) > 1:
			print len(matches)
			raise Exception('More than one matches.')

		extraction = matches[0].encode('utf-8')
		if extraction.find('codes:') > 0:
			JEL = extraction[12:-1]
		elif extraction.find('code:') > 0:
			JEL = extraction[11:-1]
		else:
			JEL = extraction[4:-1]
			if JEL[0] == ':':
				JEL = JEL[1:]

		JEL = JEL.split()
		codes = ''.join(JEL)
		codes = codes.replace('and', ',') # 'and'
		codes = codes.replace(',','\t') # ','
		codes = codes.replace(';','\t') # ';'

		break

	line = '%s\t%s\t%s\t%s\n' % (year, title, journal, codes)
	result.write(line)

	print '%s completed' %(f)

result.close()
