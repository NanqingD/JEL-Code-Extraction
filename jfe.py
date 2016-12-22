''' Python code run for JFE folder
'''


import os
import docx
import re

cur = os.getcwd()  # current working directory

filename = 'Result_JFE.csv'

try:
    os.remove(filename)
except OSError:
    pass

result = open(filename, 'wb')

jfe = cur + '\\JFE'  
os.chdir(jfe)

ncur = os.getcwd()  # new current working directory

files = os.listdir('.') # files under current working directory

journal = 'Journal of Financial Economics'


# header
header = "Year,Title,Journal"
for i in range(1,16):
	header = header + ',JEL Code ' + str(i)
header = header + '\n'
result.write(header)

for f in files:
	f = f.encode('utf-8')
	name = f[:-5]
	j = name.find('Journal')
	if j < 0:
		title = name[:-5].strip()
		year = name[-4:]
	else:
		title = name[:j-6].strip()
		year = name[j-5:j-1]

	year = int(year)

	doc = docx.Document(f)

	JEL = 'No JEL'
	pattern = re.compile('J[eE][lL]')
	for i in range(0,len(doc.paragraphs)):
		text = doc.paragraphs[i].text.encode('utf-8')
		if not re.match(pattern, text):
			continue
		extraction = text
		index = extraction.find(':')
		JEL = extraction[index+1:].strip()
		if len(JEL) > 2:
			JEL = extraction[index+1:]
			JEL = JEL.strip()
			JEL = JEL.replace(';',',')
			if JEL.find(',') >= 0:
				JEL = JEL.replace(' ','')
			else:
				codes = JEL.split()
				JEL = ','.join(codes)
		else:
			JEL = doc.paragraphs[i+1].text.encode('utf-8')
			JEL = JEL.strip()
			JEL = JEL.replace(';',',')
			if JEL.find(',') >= 0:
				JEL = JEL.replace(' ','')
			else:
				codes = JEL.split()
				JEL = ','.join(codes)
		print JEL


	# if (year <= 2007):
	# pattern = re.compile('J[eE][lL]\s+[cC]lassifications*:.+')
	# 	for i in range(0, len(doc.paragraphs)):
	# 		text = doc.paragraphs[i].text.encode('utf-8')
	# 		matches = pattern.findall(text)
	# 		if len(matches) == 0:
	# 			continue
	# 		extraction = matches[0].encode('utf-8')
	# 		index = extraction.find(':')
	# 		JEL = extraction[index+1:]
	# 		JEL = JEL.strip()
	# 		JEL = JEL.replace(';',',')
	# 		JEL = JEL.replace(' ','')
	# 		print JEL
	# 		break
	# elif (year == 2008):

	# else:
	# 	pattern = re.compile('J[eE][lL]')
	# 	for i in range(0, len(doc.paragraphs)):
	# 		text = doc.paragraphs[i].text.encode('utf-8')
	# 		if not re.match(pattern,text):
	# 			continue
	# 		JEL = doc.paragraphs[i+1].text.encode('utf-8').strip()
	# 		JEL = doc.paragraphs[i].text.encode('utf-8').strip()
	# 		index = JEL.find(':')
	# 		JEL = JEL[index+1:].strip()
	# 		JEL = JEL.replace(';',',')
	# 		JEL = JEL.replace(' ','')
	# 		print JEL
	# 		break
	

	# 	JEL = matches[0][5:-1].encode('utf-8')
	# 	JEL = JEL.split()
	# 	codes = ''.join(JEL)
	# 	codes = codes.replace(',','\t')
	# 	break

	line = '%s,%s,%s,%s\n' % (year, title, journal,JEL)
	print '%s completed'% (title)
	# print '%s' % (year)
	result.write(line)

result.close()